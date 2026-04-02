"""Tests for RFP parsers."""

from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest

from blackbox.models import ContractType, RFPContext
from blackbox.parsers import ParseError, extract_text_from_docx, extract_text_from_pdf, parse_rfp


@pytest.fixture
def sample_pdf(tmp_path) -> Path:
    """Create a small test PDF using PyMuPDF."""
    import fitz
    path = tmp_path / "test_rfp.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "REQUEST FOR PROPOSAL\n"
        "Agency: Texas Department of Information Resources\n"
        "Solicitation: DIR-TSO-TMP-456\n"
        "Scope: IT Modernization Services\n"
        "Due Date: 2026-05-01\n"
        "Contract Type: IDIQ\n",
    )
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    """Create a small test DOCX using python-docx."""
    import docx
    path = tmp_path / "test_rfp.docx"
    doc = docx.Document()
    doc.add_paragraph("REQUEST FOR PROPOSAL")
    doc.add_paragraph("Agency: FDA Center for Devices")
    doc.add_paragraph("Solicitation: FDA-RFP-2026-001")
    doc.add_paragraph("Scope: Medical Device Cybersecurity")
    doc.save(str(path))
    return path


@pytest.fixture
def empty_pdf(tmp_path) -> Path:
    """Create an empty PDF."""
    import fitz
    path = tmp_path / "empty.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.save(str(path))
    doc.close()
    return path


def _mock_claude_response(fields: dict) -> MagicMock:
    """Create a mock Claude response with tool_use."""
    tool_block = MagicMock()
    tool_block.type = "tool_use"
    tool_block.input = fields
    message = MagicMock()
    message.content = [tool_block]
    return message


class TestExtractTextFromPdf:
    def test_extracts_text(self, sample_pdf):
        text = extract_text_from_pdf(sample_pdf)
        assert "Texas Department" in text
        assert "DIR-TSO-TMP-456" in text

    def test_empty_pdf_returns_empty(self, empty_pdf):
        text = extract_text_from_pdf(empty_pdf)
        assert text == ""

    def test_invalid_file_raises(self, tmp_path):
        bad = tmp_path / "bad.pdf"
        bad.write_text("not a pdf")
        with pytest.raises(ParseError):
            extract_text_from_pdf(bad)


class TestExtractTextFromDocx:
    def test_extracts_text(self, sample_docx):
        text = extract_text_from_docx(sample_docx)
        assert "FDA Center for Devices" in text
        assert "Medical Device" in text

    def test_invalid_file_raises(self, tmp_path):
        bad = tmp_path / "bad.docx"
        bad.write_text("not a docx")
        with pytest.raises(ParseError):
            extract_text_from_docx(bad)


class TestParseRfp:
    @pytest.mark.asyncio
    async def test_parse_pdf_with_claude(self, sample_pdf):
        mock_client = AsyncMock()
        mock_client.messages.create = AsyncMock(
            return_value=_mock_claude_response({
                "agency_name": "Texas DIR",
                "scope": "IT Modernization",
                "contract_type": "IDIQ",
                "solicitation_number": "DIR-TSO-TMP-456",
            })
        )
        result = await parse_rfp(sample_pdf, mock_client, "claude-haiku-4-5-20251001")
        assert result.agency_name == "Texas DIR"
        assert result.contract_type == ContractType.IDIQ

    @pytest.mark.asyncio
    async def test_parse_docx_with_claude(self, sample_docx):
        mock_client = AsyncMock()
        mock_client.messages.create = AsyncMock(
            return_value=_mock_claude_response({
                "agency_name": "FDA",
                "scope": "Cybersecurity",
            })
        )
        result = await parse_rfp(sample_docx, mock_client, "claude-haiku-4-5-20251001")
        assert result.agency_name == "FDA"

    @pytest.mark.asyncio
    async def test_unsupported_extension_raises(self, tmp_path):
        txt = tmp_path / "file.txt"
        txt.write_text("hello")
        mock_client = AsyncMock()
        with pytest.raises(ParseError, match="Unsupported file type"):
            await parse_rfp(txt, mock_client, "model")

    @pytest.mark.asyncio
    async def test_fallback_on_claude_failure(self, sample_pdf):
        mock_client = AsyncMock()
        mock_client.messages.create = AsyncMock(side_effect=Exception("API error"))
        result = await parse_rfp(sample_pdf, mock_client, "model")
        # Should return something from heuristics, not crash
        assert isinstance(result, RFPContext)
        assert result.agency_name  # should find something

    @pytest.mark.asyncio
    async def test_empty_pdf_returns_unknown(self, empty_pdf):
        mock_client = AsyncMock()
        result = await parse_rfp(empty_pdf, mock_client, "model")
        assert result.agency_name == "Unknown Agency"
